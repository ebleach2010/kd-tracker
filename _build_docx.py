#!/usr/bin/env python3
"""Generate the iOS Home-Screen setup guide as a .docx."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# base style
st = doc.styles['Normal']
st.font.name = 'Calibri'
st.font.size = Pt(11)

def h1(t):
    p = doc.add_heading(t, level=1)
    return p

def h2(t):
    return doc.add_heading(t, level=2)

def para(t, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    return p

def bullet(t):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(t)
    return p

def numbered(t):
    p = doc.add_paragraph(style='List Number')
    p.add_run(t)
    return p

def code(t):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.font.name = 'Consolas'; r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x22, 0x55, 0x99)
    return p

# ---- Title ----
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run('NeuroTrack')
tr.bold = True; tr.font.size = Pt(26); tr.font.color.rgb = RGBColor(0xff, 0x4d, 0x9d)
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run('How to save the app to your iPhone Home Screen')
sr.italic = True; sr.font.size = Pt(13); sr.font.color.rgb = RGBColor(0x66, 0x66, 0x88)
doc.add_paragraph()

para('This guide explains, in plain steps, how a single HTML file becomes an app icon on your '
     'iPhone that opens full-screen and works offline. There are two ways to do it — Option A '
     '(a hosted link, recommended) and Option B (the file itself). No coding needed.', italic=True)

# ---- What this is ----
h1('What "an HTML app" means here')
para('NeuroTrack is one file, index.html. A web browser knows how to open it. When you "Add to '
     'Home Screen," iPhone makes an icon that opens that page like a normal app — full screen, no '
     'browser bars — and remembers it so it works without internet after the first load. All your '
     'data is stored privately on your phone; nothing is uploaded anywhere.')

# ---- Option A ----
h1('Option A — Use the hosted link (recommended)')
para('This is the easiest and always-up-to-date way. The app is published for free by GitHub Pages, '
     'which turns the file in your repository into a normal web address.')

h2('One-time setup (on a computer, ~2 minutes)')
numbered('Go to your repository on github.com (the one containing index.html).')
numbered('Click Settings (top menu of the repo).')
numbered('In the left sidebar, click Pages.')
numbered('Under "Build and deployment" → Source, choose "Deploy from a branch."')
numbered('Set Branch to main and the folder to / (root). Click Save.')
numbered('Wait about a minute, then refresh. GitHub shows a green banner with your live web '
         'address, like: https://YOURNAME.github.io/neuro-tracker/')
para('Tip: if your files are on a working branch (e.g. one starting with "claude/"), first merge '
     'that branch into main — or select that branch in step 5 — so Pages serves the latest version.',
     italic=True, size=10)

h2('Add it to your Home Screen (on the iPhone)')
numbered('Open Safari (this must be Safari, not Chrome).')
numbered('Type or paste your github.io web address and load the page.')
numbered('Tap the Share button (the square with an up-arrow, at the bottom of the screen).')
numbered('Scroll down and tap "Add to Home Screen."')
numbered('Give it a name (e.g. NeuroTrack) and tap Add.')
numbered('Close Safari and tap the new icon. It opens full-screen. Do the assessment once, then '
         'it will also work with no internet.')

# ---- Option B ----
h1('Option B — Use the file directly (no website)')
para('If you would rather not publish anything, you can carry the file itself. This works, but '
     'updating it later is more manual, and offline behavior is slightly less reliable than Option A.')
numbered('Get index.html onto your iPhone — e.g. email it to yourself, or save it to the Files app '
         '(iCloud Drive / On My iPhone).')
numbered('Open the Files app and tap index.html. It should open in a preview.')
numbered('If it opens as a web page with a Share button, tap Share → "Add to Home Screen."')
para('Note: iOS sometimes previews HTML instead of opening it as a real page. If "Add to Home '
     'Screen" is not offered, use Option A instead — hosting is the reliable path on iPhone.',
     italic=True, size=10)

# ---- Updating ----
h1('Updating the app later')
para('Option A (hosted): whenever the file in your repository changes, GitHub Pages updates the '
     'live link automatically. On the phone, open the app; if it looks stale, open the link once in '
     'Safari and pull down to refresh, then reopen the Home-Screen icon.')
para('Option B (file): replace the old index.html with the new one and add it to the Home Screen '
     'again.')

# ---- Your data ----
h1('About your data (important)')
bullet('Everything you log stays on your phone (in the browser\'s local storage). It is private and '
       'never sent anywhere.')
bullet('Because it lives in the browser, iOS can occasionally clear it if storage runs low or you '
       'clear Safari data. So back up regularly.')
bullet('To back up: open the app → Settings → Data → "Backup JSON." Keep the file somewhere safe '
       '(Files app, email). To move to a new phone or recover, use "Restore JSON."')
bullet('You can also export your history as a spreadsheet: Settings → Data → "Export CSV."')

# ---- Reminders ----
h1('Getting timed reminders')
para('A Home-Screen web app can show reminders when it is open, but iPhone will not reliably send '
     'timed pop-ups in the background. For dependable alarms (morning mood + 8am glucose, the 8pm '
     'check-in, the Thursday weekly, glucose and blood-pressure times):')
numbered('Open the app → Settings → Reminders.')
numbered('Tap "Export reminders (.ics)."')
numbered('Open the downloaded file and add the events to your Apple Calendar. Calendar will then '
         'give you real alerts at those times.')

# ---- Quick reference ----
h1('Quick reference')
para('Daily (evening): mood, sliders, meds/supplements (log only changes; mg, 0 = skipped), meals, '
     'self-care, vitals, quick note tags.')
para('Morning: tap "Morning mood" (and log wake glucose).')
para('Weekly (Thursday evening): the full version adds the cognitive tests and questionnaires.')
para('The pink Neuroinflammation line is a symptom-based estimate, not a lab test. It is for '
     'spotting your own trends — not for diagnosing a relapse. For medical decisions, contact your '
     'neurologist.', italic=True)

doc.save('NeuroTrack-iPhone-Setup-Guide.docx')
print('wrote NeuroTrack-iPhone-Setup-Guide.docx')
